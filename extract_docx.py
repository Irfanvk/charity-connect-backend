from docx import Document
import os

file_path = 'charity_management_project_plan.docx'

try:
    doc = Document(file_path)
    
    print("=" * 80)
    print("PROJECT PLAN CONTENT")
    print("=" * 80)
    print()
    
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n" + "=" * 80)
    print("TABLES")
    print("=" * 80)
    
    for table in doc.tables:
        for row in table.rows:
            print(" | ".join(cell.text for cell in row.cells))
        print()
    
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
